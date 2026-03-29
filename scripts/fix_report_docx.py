# -*- coding: utf-8 -*-
"""
fix_report_docx.py
批量修正诊断图总结报告中的 FDR 值和叙事文字。
"""

from docx import Document
import re

INPUT_DOCX = r"C:\Users\Lenovo\Desktop\诊断图总结报告.docx"
OUTPUT_DOCX = r"C:\Users\Lenovo\Desktop\诊断图总结报告_修正版.docx"

# ============================================================
# 1. 简单的文本替换规则 (旧文本 → 新文本)
# ============================================================
SIMPLE_REPLACEMENTS = [
    # --- 第一层 背景场 corr FDR ---
    ("FDR 显著比例 = 29.6% ⬆⬆",
     "FDR 显著比例 = 47.3% ⬆⬆⬆（1636/3456）"),
    ("FDR 显著比例 = 14.3% ⬆",
     "FDR 显著比例 = 44.2% ⬆⬆⬆（1529/3456）"),
    ("FDR 显著比例 = 3.5% ❌（几乎不显著）",
     "FDR 显著比例 = 20.5%（714/3480）⬆⬆"),
    ("FDR 显著比例 = 2.0% ❌",
     "FDR 显著比例 = 13.4%（467/3480）⬆"),

    # --- 第二层 背景场 diff FDR ---
    ("FDR 显著格点占比 = 38.5% ⬆⬆⬆",
     "FDR 显著格点占比 = 31.9%（1103/3456）⬆⬆⬆"),
    ("FDR 显著格点占比 = 23.1% ⬆⬆",
     "FDR 显著格点占比 = 27.0%（933/3456）⬆⬆"),
    ("FDR 显著格点占比 = 6.3% ❌",
     "FDR 显著格点占比 = 0.0%（0/3456）❌"),
    ("FDR 显著格点占比 = 8.5% ❌",
     "FDR 显著格点占比 = 18.1%（626/3456）⬆"),
    ("FDR 显著格点占比 = 6.1% ❌",
     "FDR 显著格点占比 = 15.9%（548/3456）⬆"),

    # --- 第一层总结 ---
    ("背景场：u 的 FDR = 29.6%，T 的 FDR = 14.3%——信号清晰",
     "背景场：u 的 FDR = 47.3%，T 的 FDR = 44.2%，q = 32.1%，v = 38.2%，w = 31.8%——全部变量均有显著信号"),
    ("扰动场：5 个变量的 FDR 全部 ≤ 3.5%——基本无信号",
     "扰动场：信号较弱但并非全无——u = 20.5%，v = 26.6%，T = 13.4%，q = 3.2%，w = 3.2%"),

    # --- 柱积分 ---
    ("FDR 显著比例 = 57.6% ⬆⬆⬆⬆⬆",
     "FDR 显著比例 = 51.4%（74/144）⬆⬆⬆⬆⬆"),

    # --- 水汽复活叙事 ---
    ("重要性的飞跃：回忆第一层 bg_q_corr.png 的 FDR = 0.0%——2D 场中水汽完全无信号！柱积分后信号从 0.0% 暴涨到 57.6%，证明垂直积分消除了 2D 中的高低层信号抵消，揭示了一个此前隐藏的超强信号。",
     "信号的增强：第一层 bg_q_corr.png 的 FDR = 32.1%——2D 场中水汽已有分散的信号。柱积分后进一步集中为 51.4%，表明垂直积分有效消除了高低层的部分抵消，使水汽的净信号更加清晰和集中。"),

    ("38.5% 是第一、二层中所有 diff 图的最高 FDR 值。",
     "u 的 diff FDR（31.9%）是背景场中最高的 diff 值之一。"),

    # --- 第三层总结 ---
    ("水汽信号的\"复活\"：2D 场中 q 的 FDR = 0.0%（无信号），柱积分后 Col q 的 FDR = 57.6%（最强信号）。这是本层最重要的方法学发现——垂直积分消除了 2D 中高低层的信号抵消。",
     "水汽信号的\"增强\"：2D 场中 q 的 FDR = 32.1%（分散信号），柱积分后 Col q 的 FDR = 51.4%（集中信号）。垂直积分有效消除了高低层的部分抵消，使水汽的净效应更加清晰。"),

    # --- Col q 方向修正 ---
    ("暖池区（80°–160°E）几乎全部经度上都有强负相关（蓝色）",
     "大部分经度上呈正相关（红色显著），仅暖池核心（100°–130°E）附近有窄带负相关区域"),
    ("r 值在暖池中心（120°E）达到 −0.35 以上",
     "正相关 r 值在多个经度上达到 0.25–0.31"),
    ("物理解读：背景柱水汽越高 → MJO 速度越慢。这一方向看似反直觉（水汽是\"燃料\"，多水汽不是应该加速吗？），但有深刻的物理机制：",
     "物理解读：整体上背景柱水汽与 MJO 速度呈正相关（水汽越多→速度越快），但暖池核心存在局部负相关区。正相关的物理机制较为直观："),
    ("暖池区背景水汽已经很高时，大气接近饱和，潜在不稳定度下降，垂直运动效率降低",
     "水汽充足的环境为 MJO 前方的对流触发提供了更好的条件——\"moisture mode\" 理论的基本预测"),
    ("\"Gross Moist Stability\"理论：当环境水汽过高时，上升运动输出的 MSE 增加，MJO 的有效能量辐散增强，削弱传播",
     "暖池核心附近的局部负相关可能反映了 \"Gross Moist Stability\" 效应：当环境水汽过高时，大气接近饱和，垂直运动效率降低"),
    ("这可能也反映了 ENSO 调制——La Niña 期间暖池区偏湿，同时 MJO 活动可能减慢",
     "此外，ENSO 的调制作用也不可忽视——不同 ENSO 位相下暖池水汽的分布和 MJO 速度均有系统性差异"),

    # --- Col q 小结修正 ---
    ("核心发现：背景柱积分水汽是全部分析中最强的 1D 预测因子！暖池区背景越湿 → MJO 越慢。扰动坐标下信号仍弱，再次确认\"环境 > 内部结构\"。",
     "核心发现：背景柱积分水汽是最强的 1D 预测因子之一。整体上呈正相关（水汽多→速度快），暖池核心有局部负相关。扰动坐标下信号仍弱，再次确认\"环境 > 内部结构\"。"),

    # --- LHF 方向修正 ---
    ("暖池东侧（120°–160°E）正相关：背景 LHF 更强时 MJO 更快",
     "暖池区（40°–80°E 及 130°–170°E）显著负相关：背景 LHF 更强时 MJO 速度更慢"),
    ("显著格点集中在暖池东边界一带",
     "赤道非洲附近（-10°–10°E）有窄带正相关。显著格点分散在多个区域"),
    ("物理解读：海面蒸发更强 → 低层大气水汽供给更充足 → MJO 前方的预湿润效率更高 → 对流更易在前方触发 → 加速东传。这是海洋下边界对 MJO 的水汽强迫机制。",
     "物理解读：暖池区 LHF 偏高可能反映了更强的表面风速，但同时也可能指示更活跃的背景对流——后者可能\"锚定\" MJO 降低传播速度。LHF 与速度的负相关表明海-气耦合对 MJO 速度的影响比简单的\"蒸发→加速\"更为复杂。"),

    # --- 第一层总结修正 ---
    ("水汽的 2D 信号为何缺失？ q 在背景场和扰动场中 FDR 均为 0.0%。但 moisture mode 理论预测水汽应该重要——这个矛盾将在第三层（柱积分分析）中得到解决。",
     "水汽的 2D 信号相对分散（FDR=32.1%），不如 u（47.3%）集中。moisture mode 理论预测水汽应是核心因子——柱积分分析（第三层）将进一步评估其信号。"),
    ("MJO 速度差异由大尺度环境背景场控制，而非 MJO 自身的内部结构。",
     "MJO 速度差异主要由大尺度环境背景场控制，MJO 自身的内部结构也有一定的信号但弱于背景场。"),
    ("这一对比极为鲜明，排除了\"MJO 速度由自身结构决定\"的可能",
     "背景场信号整体强于扰动场，但扰动场并非完全无信号——特别是 u 和 v 的扰动也有一定的 FDR（分别为 20.5% 和 26.6%）"),

    # --- 第二层总结修正 ---
    ("核心发现：全部 5 个变量的 FDR 均 ≤ 2.2%。Fast 和 Slow MJO 的内部扰动结构没有任何统计显著差异。",
     "核心发现：全部 5 个变量的 diff FDR 均较低（≤ 2.2%）。Fast 和 Slow MJO 的内部扰动结构差异不如背景场显著。"),
    ("核心发现与第一层完全一致：u 差异最显著，T 次之，q/v/w 不显著。分组对比的直观展示进一步巩固了\"背景西风控制速度\"的结论。",
     "核心发现与第一层方向一致：u 和 T 的差异在 diff 中最显著（FDR 31.9%/27.0%），v 和 w 也有一定信号（18.1%/15.9%），仅 q 的 diff FDR = 0.0%。分组对比进一步巩固了\"背景场控制速度\"的结论。"),
    ("重要性：u 是背景场中唯一产生大面积显著信号的变量。这表明纬向风场是影响 MJO 速度的最重要环境因子之一。",
     "重要性：u 是背景场中 FDR 最高的变量（47.3%），且信号空间集中度最好。其他变量也有显著信号但空间更分散。纬向风场是影响 MJO 速度的最重要环境因子。"),
    ("第一层通过\"相关性\"发现背景 u 与速度有关，扰动场与速度无关。第二层通过\"分组对比\"用不同方法再次得到相同结论，形成交叉验证：",
     "第一层发现背景场全部变量均与速度有统计显著联系（u 最强），扰动场信号较弱但不完全缺失。第二层通过分组对比验证了背景场主导的结论："),
    ("速度变异由大尺度环境控制，MJO 自身结构不决定速度。",
     "速度变异主要由大尺度环境控制，MJO 自身结构的影响相对较弱。"),

    # --- 第三层遗留问题 ---
    ("Col q 负相关的机制：背景水汽更多 → 速度更慢，具体通过什么动力学过程？",
     "Col q 与速度关系的机制：柱积分水汽与速度的空间分布较为复杂（整体正相关但暖池核心局部负相关），具体通过什么动力学过程？"),

    # --- 坐标系修正 ---
    ("横轴 0°–360°E",
     "横轴 -180°–180°E"),

    # --- 第二层前瞻 ---
    ("第一、二层建立了\"背景 u 最重要\"的共识。但水汽（q）在 2D 场中始终没有信号——这与 moisture mode 理论的预期矛盾。接下来的第三层将通过柱积分和表面变量分析，彻底改变对 q 重要性的认知。",
     "第一、二层建立了\"背景场主导、u 最强\"的共识。水汽（q）在 2D 场中的信号（FDR=32.1%）相对分散。接下来的第三层通过柱积分和表面变量分析，将进一步评估水汽的净效应。"),
    ("第一、二层在经度-高度（lon × level）2D 截面上分析了 5 个大气变量，发现水汽 q 在 2D 场上 FDR = 0.0%——完全没有信号。但 moisture mode 理论预测水汽应该是控制 MJO 传播的核心因素，这产生了矛盾。",
     "第一、二层在经度-高度 2D 截面上分析了 5 个大气变量，发现水汽 q 在 2D 场上虽有信号（FDR = 32.1%）但不如纬向风 u（47.3%）集中。moisture mode 理论预测水汽应是核心因素，这提示需要进一步提炼水汽的净信号。"),
    ("关键洞见是：2D 场中水汽的信号在垂直方向上被抵消了。低层和高层的水汽-速度相关方向可能相反，逐格点分析中两者分别不显著，但把它们在垂直方向上积分（柱积分），净效应就可能是显著的。",
     "关键洞见是：2D 场中水汽的信号在垂直方向上存在部分抵消。低层和高层的水汽-速度相关方向可能不同，导致信号分散。通过垂直积分（柱积分），可以提取出更清晰的净信号。"),
]

# ============================================================
# 2. 上下文相关替换（根据前文匹配特定段落）
# ============================================================
CONTEXT_REPLACEMENTS = [
    {
        "context_before": "bg_v_corr.png",
        "old": "FDR 显著比例 = 0.0% ❌",
        "new": "FDR 显著比例 = 38.2%（1319/3456）⬆⬆⬆",
        "max_gap": 5,
    },
    {
        "context_before": "bg_v_corr.png",
        "old": "核心结果：无显著信号。符合赤道纬圈平均后经向风信号弱的预期。",
        "new": "核心结果：高层和低层多个经度有显著信号，但空间分布较分散，不如 u 集中。",
        "max_gap": 5,
    },
    {
        "context_before": "bg_w_corr.png",
        "old": "FDR 显著比例 = 0.0% ❌",
        "new": "FDR 显著比例 = 31.8%（1098/3456）⬆⬆",
        "max_gap": 5,
    },
    {
        "context_before": "bg_w_corr.png",
        "old": "核心结果：无显著信号。背景的上升/下沉运动强度与 MJO 速度无直接联系。",
        "new": "核心结果：暖池区低层有显著正相关，东太平洋有负相关。背景垂直运动与 MJO 速度存在一定联系。",
        "max_gap": 5,
    },
    {
        "context_before": "2.3 背景场小结",
        "old": "核心发现：在 2D 背景场中，纬向风 u 是唯一产生广泛显著信号的变量。暖池区低层西风越强、高层东风越弱，MJO 传播越快。",
        "new": "核心发现：在 2D 背景场中，全部 5 个变量均有统计显著信号（FDR 31.8%–47.3%），其中 u 的 FDR 最高且空间最集中。暖池区低层西风越强、高层东风越弱，MJO 传播越快。",
        "max_gap": 5,
    },
    {
        "context_before": "mjo_q_corr.png",
        "old": "FDR 显著比例 = 0.0% ❌",
        "new": "FDR 显著比例 = 3.2%（111/3480）❌（弱）",
        "max_gap": 5,
    },
    {
        "context_before": "mjo_v_corr.png",
        "old": "FDR 显著比例 = 0.0% ❌",
        "new": "FDR 显著比例 = 26.6%（927/3480）⬆⬆",
        "max_gap": 5,
    },
    {
        "context_before": "mjo_w_corr.png",
        "old": "FDR 显著比例 = 1.6% ❌",
        "new": "FDR 显著比例 = 3.2%（113/3480）❌（弱）",
        "max_gap": 5,
    },
    {
        "context_before": "3.3 扰动场小结",
        "old": "核心发现：MJO 的内部结构（对流中心附近的环流、水汽、温度分布）与相速度几乎没有统计显著联系。换言之：快速和慢速 MJO 的内部结构没有系统性差异。",
        "new": "核心发现：扰动场信号弱于背景场，但不完全缺失。u（20.5%）和 v（26.6%）有一定信号，T（13.4%）次之，q 和 w（各 3.2%）接近噪声。",
        "max_gap": 5,
    },
    {
        "context_before": "mjo_u_corr.png",
        "old": "核心结果：仅在中心东侧（+30°~+60°）6 km 处有零星正相关点。绝大部分格点不显著。",
        "new": "核心结果：高层（300–200 hPa）中心西侧有大面积正相关（FDR 通过），低层有局部负相关中心。信号弱于背景 u 但不可忽略。",
        "max_gap": 5,
    },
    # bg_q_corr 需要特殊处理（因为 FDR=0.0% 出现多次）
    {
        "context_before": "bg_q_corr.png",
        "old": "FDR 显著比例 = 0.0% ❌",
        "new": "FDR 显著比例 = 32.1%（1111/3456）⬆⬆",
        "max_gap": 5,
    },
    {
        "context_before": "bg_q_corr.png",
        "old": "核心结果：没有任何格点通过 FDR 校正。 在 2D 经度-高度截面上，背景水汽的空间分布与 MJO 速度无统计显著关系。",
        "new": "核心结果：多个层次和经度的格点通过 FDR 校正。背景水汽在 2D 场中已有一定程度的信号，但空间分布较分散，不如纬向风 u 集中。",
        "max_gap": 5,
    },
    {
        "context_before": "bg_q_corr.png",
        "old": "重要注意：这个\"无信号\"结果后来被第三层的柱积分分析完全翻转——柱积分水汽的 FDR 高达 57.6%。原因是 2D 场中高层和低层的水汽-速度关系方向相反，彼此抵消。只有做了垂直积分才能揭示隐藏的净信号。",
        "new": "注意：虽然 2D 场中水汽已有信号（32.1%），但第三层的柱积分分析将进一步提炼：柱积分水汽的 FDR 达到 51.4%，说明垂直积分消除了高低层的部分抵消，使信号更加集中。",
        "max_gap": 5,
    },
    # bg_q_diff - 与"bg_q_corr FDR = 0.0%"联动
    {
        "context_before": "bg_q_diff.png",
        "old": "物理解读：与第一层 bg_q_corr.png（FDR = 0.0%）一致——2D 场中水汽没有显著信号。但这个\"不显著\"的结论将在第三层被柱积分分析推翻。原因是低层负差值和中层正差值在垂直方向上相互抵消；只有把它们加起来（柱积分）才能看到净效应。",
        "new": "物理解读：虽然 bg_q_corr.png 的 FDR = 32.1% 显示 2D 场中有一定信号，但分组 diff 的 FDR = 0.0% 表明 Fast/Slow 两组的水汽差异在 2D 场上不显著。低层负差值和中层正差值在垂直方向上相互抵消；柱积分将揭示净效应。",
        "max_gap": 5,
    },
]


def replace_paragraph_text(para, old_text, new_text):
    """Replace text in paragraph while using first run's formatting."""
    full_text = para.text
    if old_text not in full_text:
        return False

    new_full = full_text.replace(old_text, new_text)

    # Clear all runs and set new text in first run
    if para.runs:
        for run in para.runs:
            run.text = ""
        para.runs[0].text = new_full
    else:
        para.add_run(new_full)
    return True


def main():
    print(f"正在读取: {INPUT_DOCX}")
    doc = Document(INPUT_DOCX)
    total = len(doc.paragraphs)
    print(f"段落总数: {total}")

    # 1. 简单替换
    print("\n--- 简单替换 ---")
    simple_hit = 0
    simple_miss = 0
    for old, new in SIMPLE_REPLACEMENTS:
        found = False
        for para in doc.paragraphs:
            if old in para.text:
                replace_paragraph_text(para, old, new)
                print(f"  ✅ '{old[:50]}...'")
                found = True
                simple_hit += 1
                break  # 每条规则只替换第一个匹配
        if not found:
            simple_miss += 1
            print(f"  ⚠️ 未找到: '{old[:50]}...'")

    print(f"\n简单替换: {simple_hit} 成功, {simple_miss} 未找到")

    # 2. 上下文替换
    print("\n--- 上下文替换 ---")
    ctx_hit = 0
    ctx_miss = 0
    para_texts = [p.text for p in doc.paragraphs]

    for rule in CONTEXT_REPLACEMENTS:
        ctx = rule["context_before"]
        old = rule["old"]
        new = rule["new"]
        max_gap = rule["max_gap"]
        done = False

        for i, text in enumerate(para_texts):
            if ctx in text:
                for j in range(i+1, min(i+1+max_gap+5, total)):
                    if old in doc.paragraphs[j].text:
                        replace_paragraph_text(doc.paragraphs[j], old, new)
                        # Update cached text
                        para_texts[j] = doc.paragraphs[j].text
                        print(f"  ✅ ctx='{ctx[:25]}' → P{j}: '{old[:40]}...'")
                        done = True
                        ctx_hit += 1
                        break
                if done:
                    break
        if not done:
            ctx_miss += 1
            print(f"  ⚠️ ctx='{ctx[:25]}', 未找到: '{old[:40]}...'")

    print(f"\n上下文替换: {ctx_hit} 成功, {ctx_miss} 未找到")

    # Save
    print(f"\n保存修正版到: {OUTPUT_DOCX}")
    doc.save(OUTPUT_DOCX)
    print("✅ 保存完成！")

    # Verify
    print("\n--- 验证 ---")
    doc2 = Document(OUTPUT_DOCX)
    old_patterns = [
        ("FDR = 0.0%", ["SHF", "Q_rad", "mjo_diff", "0/144"]),  # SHF/Q_rad的0.0%是正确的
        "FDR = 29.6%", "FDR = 14.3%", "FDR = 3.5%",
        "FDR 显著比例 = 0.0% ❌",
        "水汽完全无信号", "完全没有信号", "唯一产生大面积显著信号",
    ]
    problems = []
    for i, p in enumerate(doc2.paragraphs):
        text = p.text
        for pat in old_patterns:
            if isinstance(pat, tuple):
                pat_str, exceptions = pat
                if pat_str in text and not any(ex in text for ex in exceptions):
                    problems.append(f"  P{i}: '{pat_str}' → {text[:80]}...")
            elif isinstance(pat, str):
                if pat in text:
                    problems.append(f"  P{i}: '{pat}' → {text[:80]}...")

    if problems:
        print(f"⚠️ 发现 {len(problems)} 处可能残留:")
        for p in problems:
            print(p)
    else:
        print("✅ 验证通过：未发现残留的旧值")


if __name__ == "__main__":
    main()
