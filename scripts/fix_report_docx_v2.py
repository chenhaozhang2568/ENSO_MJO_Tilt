# -*- coding: utf-8 -*-
"""
fix_report_docx_v2.py - 二次修正
"""

from docx import Document

INPUT_DOCX = r"C:\Users\Lenovo\Desktop\诊断图总结报告_修正版.docx"
OUTPUT_DOCX = r"C:\Users\Lenovo\Desktop\诊断图总结报告_修正版.docx"  # 覆盖

SIMPLE_REPLACEMENTS = [
    # --- P13: 9层气压面 → 24层高度网格 ---
    ("ERA5 再分析资料，1979–2022，赤道纬圈平均（20°S–20°N），9 层气压面",
     "ERA5 再分析资料，1979–2022，赤道纬圈平均（20°S–20°N），插值到 24 层均匀高度网格（0.5–12 km）"),

    # --- P34: 0°–360°E → -180°–180° ---
    ("背景坐标（Background）：使用绝对经度（0°–360°E），反映大尺度环境状态",
     "背景坐标（Background）：使用绝对经度（-180°–180°E），反映大尺度环境状态"),

    # --- Col q pair (P439-P440): 方向修正 ---
    ("Slow 组（蓝线）在暖池区柱水汽始终高于 Fast 组（红线）",
     "Fast 组（红线）在大部分经度上柱水汽高于 Slow 组（蓝线），仅暖池核心（90°–120°E）附近 Fast 低于 Slow"),
    ("差异在 100°–150°E 最大，约 1–2 kg/m²",
     "正差值在印度洋至暖池西侧（30°–80°E）和西太平洋（130°–180°E）最为明显；暖池核心（100°E）附近有负差值"),

    # --- Col q diff (P444): FDR 修正 ---
    ("FDR 显著比例 = 51.8% ⬆⬆⬆⬆",
     "FDR 显著比例 = 57.6%（83/144）⬆⬆⬆⬆⬆"),

    # --- Col q diff (P445): 叙事修正 ---
    ("核心结果：暖池区强负差值（Fast 组水汽更低），大面积通过 FDR 校正。",
     "核心结果：大部分经度为正差值（Fast 组水汽更高），仅暖池核心（~100°E）有窄带负差值（Fast 更低）。大面积通过 FDR 校正。"),

    # --- LHF diff (P510): 方向修正 ---
    ("Fast 组暖池区 LHF 更高。",
     "Fast 组在大部分区域 LHF 更低（负差值），仅赤道非洲附近(-30°–10°E)有正差值。"),

    # --- SST 因果链 (P524): 与 LHF 负相关一致性 ---
    ("热力学路径：SST 高 → 海气温差/湿度差大 → 蒸发增强（LHF 增大）→ 水汽供给增加",
     "热力学路径：SST 高 → 海气温差大 → 有利于对流触发。但 LHF 与速度呈负相关（第三层），说明 SST 对速度的促进不是通过简单增强蒸发实现的"),

    # --- bg_v_diff (P282): 叙事与新FDR矛盾 ---
    ("核心结果：差异弱且零散。赤道纬圈平均后经向风信号被抵消。无显著 Fast/Slow 差异。",
     "核心结果：FDR=18.1% 表明有一定信号，但空间分布较分散。赤道纬圈平均后经向风差异弱于纬向风。"),

    # --- bg_w_diff (P289): 叙事与新FDR矛盾 ---
    ("核心结果：暖池区有弱上升运动差异，但不通过 FDR 校正。背景的上升运动强弱不显著影响速度。",
     "核心结果：FDR=15.9% 表明暖池区上升运动差异有一定统计显著性。背景垂直运动的强弱对速度有一定影响。"),
]

CONTEXT_REPLACEMENTS = [
    # --- 第三层 6.2 海洋下边界叙事（P620）: SST→LHF→水汽 需加注 ---
    {
        "context_before": "海洋下边界的首次登场",
        "old": "暖 SST 增强蒸发、供给水汽——这为第四层的因果链分析奠定基础。",
        "new": "暖 SST 与速度正相关，而 LHF 与速度负相关——这一看似矛盾的结果将在第四层的因果链分析中得到解释。",
        "max_gap": 3,
    },
]


def replace_paragraph_text(para, old_text, new_text):
    full_text = para.text
    if old_text not in full_text:
        return False
    new_full = full_text.replace(old_text, new_text)
    if para.runs:
        for run in para.runs:
            run.text = ""
        para.runs[0].text = new_full
    else:
        para.add_run(new_full)
    return True


def main():
    print(f"读取: {INPUT_DOCX}")
    doc = Document(INPUT_DOCX)
    total = len(doc.paragraphs)
    print(f"段落: {total}")

    # Simple
    print("\n--- 简单替换 ---")
    hit, miss = 0, 0
    for old, new in SIMPLE_REPLACEMENTS:
        found = False
        for para in doc.paragraphs:
            if old in para.text:
                replace_paragraph_text(para, old, new)
                print(f"  OK '{old[:50]}...'")
                found = True
                hit += 1
                break
        if not found:
            miss += 1
            print(f"  !! '{old[:50]}...'")
    print(f"简单: {hit} ok, {miss} miss")

    # Context
    print("\n--- 上下文替换 ---")
    para_texts = [p.text for p in doc.paragraphs]
    for rule in CONTEXT_REPLACEMENTS:
        ctx, old, new, gap = rule["context_before"], rule["old"], rule["new"], rule["max_gap"]
        done = False
        for i, t in enumerate(para_texts):
            if ctx in t:
                for j in range(i, min(i+gap+5, total)):
                    if old in doc.paragraphs[j].text:
                        replace_paragraph_text(doc.paragraphs[j], old, new)
                        para_texts[j] = doc.paragraphs[j].text
                        print(f"  OK ctx='{ctx[:20]}' P{j}")
                        done = True
                        break
                if done: break
        if not done:
            print(f"  !! ctx='{ctx[:20]}' not found")

    doc.save(OUTPUT_DOCX)
    print(f"\n保存: {OUTPUT_DOCX}")
    print("Done!")


if __name__ == "__main__":
    main()
