# -*- coding: utf-8 -*-
"""fix_report_docx_v3.py - Third round of fixes"""

from docx import Document
import json

INPUT_DOCX = r"C:\Users\Lenovo\Desktop\诊断图总结报告_修正版.docx"
OUTPUT_DOCX = r"C:\Users\Lenovo\Desktop\诊断图总结报告_修正版.docx"
RULES_JSON = r"C:\tmp\v3_rules.json"


def replace_paragraph_text(para, old_text, new_text):
    full_text = para.text
    if old_text not in full_text:
        return False
    new_full = full_text.replace(old_text, new_text)
    if para.runs:
        for run in para.runs:
            run.text = ""
        para.runs[0].text = new_full
    else:
        para.add_run(new_full)
    return True


def main():
    with open(RULES_JSON, "r", encoding="utf-8") as f:
        rules = json.load(f)

    print(f"Read: {INPUT_DOCX}")
    doc = Document(INPUT_DOCX)
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Rules: {len(rules)}")

    hit, miss = 0, 0
    for old, new in rules:
        found = False
        for para in doc.paragraphs:
            if old in para.text:
                replace_paragraph_text(para, old, new)
                tag = old[:60].replace("\n", " ")
                print(f"  OK '{tag}...'")
                found = True
                hit += 1
                break
        if not found:
            miss += 1
            tag = old[:60].replace("\n", " ")
            print(f"  !! MISS '{tag}...'")
    print(f"\nResult: {hit} ok, {miss} miss")

    doc.save(OUTPUT_DOCX)
    print(f"Saved: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
